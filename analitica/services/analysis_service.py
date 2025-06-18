# analitica/services/analysis_service.py
import pandas as pd
from typing import List, Dict, Any, Optional
import matplotlib
matplotlib.use('Agg') # Ensure backend is set before pyplot import
import matplotlib.pyplot as plt
import seaborn as sns
import io
from docx import Document
from docx.shared import Pt # Inches might not be used, Pt can be for font sizes
from docx.enum.text import WD_ALIGN_PARAGRAPH # Corrected import
from datetime import datetime

# Import the new DataFrame getters
from ..database import (
    get_hardware_df,
    get_employees_df,
    get_licenses_df,
    get_web_accesses_df
)
from ..schemas import CountByItem # Still useful

def get_hardware_counts_by_type_pandas() -> List[CountByItem]:
    hardware_df = get_hardware_df()
    if hardware_df is None or hardware_df.empty:
        return []

    # Ensure 'type' column exists
    if 'type' not in hardware_df.columns:
        # If type column is missing, perhaps return empty or handle error
        # For now, let's assume it might be missing and treat as 'N/A'
        # Or, if critical, return an empty list or raise error
        # Given the original code filledna('N/A'), let's try to replicate
        # If 'type' column is missing, all will be N/A
        counts_df = pd.DataFrame({'type': ['N/A'], 'id': [len(hardware_df)]})
    else:
        # Fill NaN in 'type' column if it exists
        hardware_df['type'] = hardware_df['type'].fillna('N/A')
        counts_df = hardware_df.groupby('type')['id'].count().reset_index(name='count')

    result_list = [CountByItem(item=row['type'], count=row['count']) for index, row in counts_df.iterrows()]
    return result_list

# generate_hardware_type_distribution_chart remains the same as it doesn't use `db` directly
def generate_hardware_type_distribution_chart(data: List[CountByItem]) -> io.BytesIO:
    if not data:
        fig, ax = plt.subplots()
        ax.text(0.5, 0.5, "No data available for chart", ha='center', va='center', fontsize=12)
        ax.set_xticks([]); ax.set_yticks([])
        buf = io.BytesIO(); fig.savefig(buf, format='png'); plt.close(fig); buf.seek(0)
        return buf
    items = [d.item for d in data]; counts = [d.count for d in data]
    fig, ax = plt.subplots(figsize=(10, max(6, len(items) * 0.5))) # Dynamic height
    sns.barplot(x=counts, y=items, ax=ax, palette="viridis", orient='h')
    ax.set_title('Hardware Distribution by Type', fontsize=16)
    ax.set_xlabel('Count', fontsize=12); ax.set_ylabel('Hardware Type', fontsize=12)
    for i, v in enumerate(counts): ax.text(v + 0.5, i, str(v), color='blue', va='center', fontweight='bold')
    plt.tight_layout(); buf = io.BytesIO(); fig.savefig(buf, format='png'); plt.close(fig); buf.seek(0)
    return buf

def generate_employee_acta_word(employee_id: int) -> Optional[io.BytesIO]:
    employees_df = get_employees_df()
    hardware_df = get_hardware_df()
    licenses_df = get_licenses_df()
    web_accesses_df = get_web_accesses_df()

    if employees_df is None:
        return None # Or raise error

    employee_series = employees_df[employees_df['id'] == employee_id].iloc[0] if not employees_df[employees_df['id'] == employee_id].empty else None
    if employee_series is None:
        return None

    # Filter assets for the given employee
    employee_hardware = hardware_df[hardware_df['employee_id'] == employee_id] if hardware_df is not None else pd.DataFrame()
    employee_licenses = licenses_df[licenses_df['employee_id'] == employee_id] if licenses_df is not None else pd.DataFrame()
    employee_web_accesses = web_accesses_df[web_accesses_df['employee_id'] == employee_id] if web_accesses_df is not None else pd.DataFrame()

    document = Document()
    title = document.add_heading('Acta de Asignación de Activos TIC', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()
    document.add_heading('Información del Empleado', level=2)
    document.add_paragraph(f"Nombre: {employee_series.get('name', 'N/A')}")
    document.add_paragraph(f"Departamento: {employee_series.get('department', 'N/A')}")
    document.add_paragraph(f"Cargo: {employee_series.get('position', 'N/A')}")
    document.add_paragraph(f"Fecha de Generación: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    document.add_paragraph()

    document.add_heading('Equipos de Hardware Asignados', level=2)
    if not employee_hardware.empty:
        table = document.add_table(rows=1, cols=4); table.style = 'TableGrid'
        hdr_cells = table.rows[0].cells; hdr_cells[0].text = 'Tipo'; hdr_cells[1].text = 'Marca'; hdr_cells[2].text = 'Serial'; hdr_cells[3].text = 'Ubicación'
        for _, item in employee_hardware.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = item.get('type', '')
            row_cells[1].text = item.get('brand', '')
            row_cells[2].text = item.get('serial_number', '')
            row_cells[3].text = item.get('location', '')
    else:
        document.add_paragraph("No hay equipos de hardware asignados.")
    document.add_paragraph()

    document.add_heading('Licencias de Software Asignadas', level=2)
    if not employee_licenses.empty:
        table = document.add_table(rows=1, cols=3); table.style = 'TableGrid'
        hdr_cells = table.rows[0].cells; hdr_cells[0].text = 'Software'; hdr_cells[1].text = 'Clave'; hdr_cells[2].text = 'Expiración'
        for _, item in employee_licenses.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = item.get('software_name', '')
            row_cells[1].text = item.get('license_key', '')
            exp_date = item.get('expiration_date')
            row_cells[2].text = pd.to_datetime(exp_date).strftime('%Y-%m-%d') if pd.notna(exp_date) else 'N/A'
    else:
        document.add_paragraph("No hay licencias de software asignadas.")
    document.add_paragraph()

    document.add_heading('Accesos Web Asignados', level=2)
    if not employee_web_accesses.empty:
        table = document.add_table(rows=1, cols=3); table.style = 'TableGrid'
        hdr_cells = table.rows[0].cells; hdr_cells[0].text = 'Servicio'; hdr_cells[1].text = 'URL'; hdr_cells[2].text = 'Usuario'
        for _, item in employee_web_accesses.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = item.get('service_name', '')
            row_cells[1].text = item.get('url', '')
            row_cells[2].text = item.get('access_username', '') # Column name from CSV
    else:
        document.add_paragraph("No hay accesos web asignados.")

    document.add_paragraph()
    document.add_page_break()
    document.add_heading('Firmas', level=2)
    document.add_paragraph(f"\n\n_________________________\nFirma del Empleado: {employee_series.get('name', 'N/A')}")
    document.add_paragraph("\n\n_________________________\nFirma del Responsable (Admin)")

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

def generate_full_excel_report() -> io.BytesIO:
    employees_df = get_employees_df()
    hardware_df = get_hardware_df()
    licenses_df = get_licenses_df()
    web_accesses_df = get_web_accesses_df()

    # Handle cases where DFs might be None (e.g., CSV not found)
    if employees_df is None: employees_df = pd.DataFrame()
    if hardware_df is None: hardware_df = pd.DataFrame()
    if licenses_df is None: licenses_df = pd.DataFrame()
    if web_accesses_df is None: web_accesses_df = pd.DataFrame()

    # Prepare data for Excel, similar to original but handling potential missing columns and joining for names
    # Employees
    df_employees_report = employees_df[['id', 'name', 'department', 'position']].copy()
    df_employees_report.columns = ["ID", "Name", "Department", "Position"]

    # Hardware
    df_hardware_report = hardware_df[['id', 'type', 'brand', 'serial_number', 'location', 'employee_id']].copy()
    df_hardware_report.columns = ["ID", "Type", "Brand", "Serial Number", "Location", "Assigned Employee ID"]
    if not employees_df.empty and 'id' in employees_df.columns and 'name' in employees_df.columns:
        df_hardware_report = pd.merge(df_hardware_report, employees_df[['id', 'name']], left_on='Assigned Employee ID', right_on='id', how='left', suffixes=('', '_employee'))
        df_hardware_report.rename(columns={'name': 'Assigned Employee Name'}, inplace=True)
        df_hardware_report.drop(columns=['id_employee'], errors='ignore', inplace=True) # clean up merge key if exists
    else:
        df_hardware_report['Assigned Employee Name'] = None


    # Licenses
    df_licenses_report = licenses_df[['id', 'software_name', 'license_key', 'purchase_date', 'expiration_date', 'employee_id']].copy()
    df_licenses_report.columns = ["ID", "Software Name", "License Key", "Purchase Date", "Expiration Date", "Assigned Employee ID"]
    if not employees_df.empty and 'id' in employees_df.columns and 'name' in employees_df.columns:
        df_licenses_report = pd.merge(df_licenses_report, employees_df[['id', 'name']], left_on='Assigned Employee ID', right_on='id', how='left', suffixes=('', '_employee'))
        df_licenses_report.rename(columns={'name': 'Assigned Employee Name'}, inplace=True)
        df_licenses_report.drop(columns=['id_employee'], errors='ignore', inplace=True)
    else:
        df_licenses_report['Assigned Employee Name'] = None
    # Format dates if they exist
    if 'Purchase Date' in df_licenses_report.columns:
      df_licenses_report['Purchase Date'] = pd.to_datetime(df_licenses_report['Purchase Date'], errors='coerce').dt.strftime('%Y-%m-%d')
    if 'Expiration Date' in df_licenses_report.columns:
      df_licenses_report['Expiration Date'] = pd.to_datetime(df_licenses_report['Expiration Date'], errors='coerce').dt.strftime('%Y-%m-%d')


    # Web Accesses
    df_web_accesses_report = web_accesses_df[['id', 'service_name', 'url', 'access_username', 'employee_id']].copy()
    df_web_accesses_report.columns = ["ID", "Service Name", "URL", "Access Username", "Assigned Employee ID"]
    if not employees_df.empty and 'id' in employees_df.columns and 'name' in employees_df.columns:
        df_web_accesses_report = pd.merge(df_web_accesses_report, employees_df[['id', 'name']], left_on='Assigned Employee ID', right_on='id', how='left', suffixes=('', '_employee'))
        df_web_accesses_report.rename(columns={'name': 'Assigned Employee Name'}, inplace=True)
        df_web_accesses_report.drop(columns=['id_employee'], errors='ignore', inplace=True)
    else:
        df_web_accesses_report['Assigned Employee Name'] = None

    output_stream = io.BytesIO()
    with pd.ExcelWriter(output_stream, engine='openpyxl') as writer:
        df_employees_report.to_excel(writer, sheet_name='Employees', index=False)
        df_hardware_report.to_excel(writer, sheet_name='Hardware', index=False)
        df_licenses_report.to_excel(writer, sheet_name='Licenses', index=False)
        df_web_accesses_report.to_excel(writer, sheet_name='Web Accesses', index=False)
    output_stream.seek(0)
    return output_stream

def get_employee_counts_by_department() -> List[CountByItem]:
    employees_df = get_employees_df()
    if employees_df is None or employees_df.empty:
        return []

    # Ensure 'department' column exists
    if 'department' not in employees_df.columns:
        # If 'department' column is missing, count all as 'N/A' or handle as error
        return [CountByItem(item="N/A", count=len(employees_df))]

    # Fill NaN in 'department' column if it exists
    employees_df['department'] = employees_df['department'].fillna('N/A')
    counts_df = employees_df.groupby('department')['id'].count().reset_index(name='count')

    result_list = [CountByItem(item=row['department'], count=row['count']) for index, row in counts_df.iterrows()]
    return result_list

def get_license_counts_by_software() -> List[CountByItem]:
    licenses_df = get_licenses_df()
    if licenses_df is None or licenses_df.empty:
        return []

    # Ensure 'software_name' column exists
    if 'software_name' not in licenses_df.columns:
         return [CountByItem(item="N/A", count=len(licenses_df))]

    licenses_df['software_name'] = licenses_df['software_name'].fillna('N/A')
    counts_df = licenses_df.groupby('software_name')['id'].count().reset_index(name='count')

    result_list = [CountByItem(item=row['software_name'], count=row['count']) for index, row in counts_df.iterrows()]
    return result_list
